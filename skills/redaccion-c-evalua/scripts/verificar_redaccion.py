#!/usr/bin/env python3
"""
verificar_redaccion.py — Verificación y formateo de estilo C-evalua para Producto 1 FASP.

Aplica las reglas de la guía de estilo de C-evalua a un documento Word del
Producto 1. Tiene dos modos:

    --check   Solo informa los problemas encontrados (no modifica el archivo)
    --fix     Aplica las correcciones automáticas y guarda el archivo
    --report  Genera un reporte de verificación en texto plano

Uso:

    # Solo verificar
    python3 verificar_redaccion.py --check Producto1_Queretaro_2026.docx

    # Verificar y aplicar fix
    python3 verificar_redaccion.py --fix Producto1_Queretaro_2026.docx

    # Reporte completo
    python3 verificar_redaccion.py --report Producto1_Queretaro_2026.docx

El script distingue entre problemas AUTOMÁTICOS (que puede corregir solo) y
problemas MANUALES (que requieren intervención humana). Los automáticos se
corrigen en modo --fix; los manuales solo se reportan.
"""

import argparse
import pathlib
import re
import sys
from collections import defaultdict
from typing import NamedTuple

# ─────────────────────────────────────────────────────────────────
# CONFIGURACIÓN
# ─────────────────────────────────────────────────────────────────

PYTHON3 = "/Users/adominguezdia/.hermes/hermes-agent/venv/bin/python3"
SKILL_REFS = pathlib.Path(__file__).parent.parent / "references" / "reglas_redaccion_fasp.md"


class Hallazgo(NamedTuple):
    """Un problema encontrado en el documento."""
    tipo: str          # 'AUTO' o 'MANUAL'
    severidad: str      # 'ERROR', 'WARNING', 'INFO'
    ubicacion: str     # 'párrafo N', 'tabla N', 'sección X'
    regla: str         # Identificador de la regla violada
    descripcion: str   # Texto legible del problema
    texto_original: str # El texto que tiene el problema
    sugerencia: str    # Qué se recomienda hacer (o qué se aplicó en --fix)


# ─────────────────────────────────────────────────────────────────
# REGLAS DE VERIFICACIÓN
# ─────────────────────────────────────────────────────────────────

# Frases de relleno que nunca deben aparecer
FRASES_RELLENO = [
    re.compile(r"cabe destacar que", re.IGNORECASE),
    re.compile(r"es importante mencionar que", re.IGNORECASE),
    re.compile(r"en términos generales", re.IGNORECASE),
    re.compile(r"de conformidad con lo establecido", re.IGNORECASE),
    re.compile(r"en el marco del presente", re.IGNORECASE),
    re.compile(r"a decir de los resultados", re.IGNORECASE),
    re.compile(r"los resultados arrojan que", re.IGNORECASE),
    re.compile(r"se puede observar que", re.IGNORECASE),
    re.compile(r"se evidencia que", re.IGNORECASE),
    re.compile(r"se hace necesario señalar", re.IGNORECASE),
]

# Primera persona individual (nunca debe aparecer)
PRIMERA_PERSONA = [
    re.compile(r"\bconsidero\b", re.IGNORECASE),
    re.compile(r"\bcreo que\b", re.IGNORECASE),
    re.compile(r"\bopino que\b", re.IGNORECASE),
    re.compile(r"\bpienso que\b", re.IGNORECASE),
    re.compile(r"\bme parece\b", re.IGNORECASE),
    re.compile(r"\ben mi opinión\b", re.IGNORECASE),
    re.compile(r"\bdesde mi punto de vista\b", re.IGNORECASE),
]

# Absolutos sin dato (advertencia, no error)
ABSOLUTOS_SIN_DATO = [
    re.compile(r"\btodos los\b", re.IGNORECASE),
    re.compile(r"\bningún\b", re.IGNORECASE),
    re.compile(r"\bsiempre\b", re.IGNORECASE),
    re.compile(r"\bnunca\b", re.IGNORECASE),
    re.compile(r"\b100%\b", re.IGNORECASE),
]

# Formato correcto de página: 'p. N' (no 'pág. N', 'pag. N', 'pág N')
PAGINA_INCORRECTO = re.compile(r"\bp[a]?g\.?\s+\d+", re.IGNORECASE)

# Artículo bis incorrecto: '28.5', '28BIS', '28 BIS', '28 bis'
ART_BIS_INCORRECTO = re.compile(
    r"\b(Art\.?\s*\d+[\. ]+(BIS|Bis|BIS\b))|"
    r"\b(Art\.?\s*\d+\.(\d+)(?!\s+Bis))|"
    r"\b(Art\.?\s*\d+\s+BIS(?!\s))",
    re.IGNORECASE
)
# Captura: 'Art. 28.5' -> debe ser 'Art. 28 Bis'
ART_BIS_CORRECTO = re.compile(r"\b(\d+)\.(\d+)\b")


# Siglas que deben definirse al menos una vez en el documento
SIGLAS_ESPERADAS = [
    ("FASP", "Fondo de Aportaciones para la Seguridad Pública"),
    ("SESNSP", "Sistema Nacional de Seguridad Pública"),
    ("DGVS", "Dirección General de Vinculación y Seguimiento"),
    ("SEE", "Secretaría Ejecutiva del Estado"),
    ("OIC", "Órgano Interno de Control"),
    ("ASF", "Auditoría Superior de la Federación"),
    ("SHCP", "Secretaría de Hacienda y Crédito Público"),
    ("CONEVAL", "Consejo Nacional de Evaluación de la Política de Desarrollo Social"),
    ("MIR", "Matriz de Indicadores para Resultados"),
    ("POA", "Programa Operativo Anual"),
    ("ROP", "Reglas de Operación"),
    ("FOFISP", "Fondo de Fortalecimiento de las Instituciones de Seguridad Pública"),
]

# Umbral de palabras por oración
UMBRAL_PALABRAS = 35


# ─────────────────────────────────────────────────────────────────
# ANÁLISIS DE TEXTO
# ─────────────────────────────────────────────────────────────────

def palabras(texto: str) -> int:
    return len(texto.split())


def verificar_oraciones_largas(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Oraciones con más de UMBRAL_PALABRAS palabras."""
    hallazgos = []
    oraciones = re.split(r"[.!?]+", texto)
    for i, oracion in enumerate(oraciones):
        oracion = oracion.strip()
        if not oracion:
            continue
        num_palabras = palabras(oracion)
        if num_palabras > UMBRAL_PALABRAS:
            hallazgos.append(Hallazgo(
                tipo="MANUAL",
                severidad="WARNING",
                ubicacion=f"{ubicacion} (oración {i+1})",
                regla="longitud_oracion",
                descripcion=f"Oración con {num_palabras} palabras (umbral: {UMBRAL_PALABRAS})",
                texto_original=oracion[:200],
                sugerencia="Dividir en 2 oraciones más cortas. Buscar subordinadas encadenadas."
            ))
    return hallazgos


def verificar_frases_relleno(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Frases de relleno que empobrecen el texto."""
    hallazgos = []
    for pattern in FRASES_RELLENO:
        for match in pattern.finditer(texto):
            hallazgos.append(Hallazgo(
                tipo="AUTO",
                severidad="WARNING",
                ubicacion=ubicacion,
                regla="frase_relleno",
                descripcion=f"Frase de relleno detectada: '{match.group()}'",
                texto_original=match.group(),
                sugerencia=f"Eliminar. Reemplazar con la información sustantiva que sigue."
            ))
    return hallazgos


def verificar_primera_persona(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Primera persona individual (prohibida)."""
    hallazgos = []
    for pattern in PRIMERA_PERSONA:
        for match in pattern.finditer(texto):
            hallazgos.append(Hallazgo(
                tipo="MANUAL",
                severidad="ERROR",
                ubicacion=ubicacion,
                regla="primera_persona",
                descripcion=f"Primera persona individual detectada: '{match.group()}'",
                texto_original=match.group(),
                sugerencia="Reescribir en voz impersonal ('se identificó que...') o primera persona plural institucional ('concluimos que...')."
            ))
    return hallazgos


def verificar_absolutos(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Absolutos sin dato que los respalde."""
    hallazgos = []
    for pattern in ABSOLUTOS_SIN_DATO:
        for match in pattern.finditer(texto):
            hallazgos.append(Hallazgo(
                tipo="MANUAL",
                severidad="WARNING",
                ubicacion=ubicacion,
                regla="absoluto_sin_dato",
                descripcion=f"Absoluto sin dato de respaldo: '{match.group()}'",
                texto_original=match.group(),
                sugerencia="Matizar: 'la mayoría de', 'en varios casos', 'se identificaron N de N'."
            ))
    return hallazgos


def verificar_paginas(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Formato de página: 'p. N' vs 'pág. N'."""
    hallazgos = []
    for match in PAGINA_INCORRECTO.finditer(texto):
        hallazgos.append(Hallazgo(
            tipo="AUTO",
            severidad="ERROR",
            ubicacion=ubicacion,
            regla="formato_pagina",
            descripcion=f"Formato de página incorrecto: '{match.group()}'",
            texto_original=match.group(),
            sugerencia="Usar 'p. N' (ej. 'p. 4'). No 'pág.', 'pag.' ni 'pág N'."
        ))
    return hallazgos


def verificar_articulos_bis(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Artículo bis en formato incorrecto (.5, BIS, etc.)."""
    hallazgos = []
    # Buscar el patrón de artículo bis INCORRECTO: "Art. 28.5" (decimal)
    # NO marcar "Art. 28 Bis" como error (eso es correcto)
    patron_incorrecto = re.compile(
        r"\bArt\.?\s*(\d+)\.(\d+)\b(?!.*\bBis\b)", re.IGNORECASE
    )
    for match in patron_incorrecto.finditer(texto):
        art_num = match.group(1)
        hallazgos.append(Hallazgo(
            tipo="AUTO",
            severidad="ERROR",
            ubicacion=ubicacion,
            regla="articulo_bis",
            descripcion=f"Formato de artículo bis incorrecto: '{match.group()}'",
            texto_original=match.group(),
            sugerencia=f"Debe ser 'Art. {art_num} Bis'"
        ))
    return hallazgos


def verificar_siglas(texto: str) -> tuple[set[str], list[Hallazgo]]:
    """Verifica que las siglas esperadas estén definidas (primera mención con nombre completo)."""
    definidas = set()
    no_definidas = set()

    for sigla, nombre_completo in SIGLAS_ESPERADAS:
        # ¿Aparece la sigla sin definición?
        patron_sin_def = re.compile(
            rf"\b{sigla}\b(?!\s+\([A-Z])", re.IGNORECASE
        )
        if patron_sin_def.search(texto):
            # ¿Aparece la sigla CON definición en algún otro lugar?
            patron_con_def = re.compile(
                rf"\b{sigla}\s*\([A-Z]", re.IGNORECASE
            )
            if not patron_con_def.search(texto):
                no_definidas.add(sigla)
            else:
                definidas.add(sigla)
        else:
            # La sigla no aparece en absoluto (no es error, solo no se usa)
            pass

    hallazgos = []
    for sigla in no_definidas:
        _, nombre = next((s, n) for s, n in SIGLAS_ESPERADAS if s == sigla)
        hallazgos.append(Hallazgo(
            tipo="MANUAL",
            severidad="WARNING",
            ubicacion="documento",
            regla="sigla_no_definida",
            descripcion=f"La sigla '{sigla}' aparece sin definición previa",
            texto_original=sigla,
            sugerencia=f"Primera mención debe ser: '{nombre} ({sigla})'"
        ))
    return definidas, hallazgos


def verificar_mayusculas_total(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Texto TODO EN MAYÚSCULAS que no sea título (Heading 1 o 2)."""
    hallazgos = []
    # Ignorar títulos de nivel 1 y 2 (esos SÍ van en mayúsculas)
    # Solo verificar párrafos normales
    if texto.isupper() and len(texto.split()) > 3:
        hallazgos.append(Hallazgo(
            tipo="MANUAL",
            severidad="WARNING",
            ubicacion=ubicacion,
            regla="mayusculas_total",
            descripcion="Texto en mayúsculas totales que no es un título",
            texto_original=texto[:100],
            sugerencia="Usar mayúsculas inicial en oraciones. Las mayúsculas totales solo en títulos de nivel 1-2."
        ))
    return hallazgos


def verificar_comillas_citas(texto: str, ubicacion: str) -> list[Hallazgo]:
    """Citas textuales sin comillas angulares."""
    hallazgos = []
    # Detectar comillas rectas "..." que no sean parte de una URL
    patron = re.compile(r'"([^"]{20,})"')
    for match in patron.finditer(texto):
        citacion = match.group(1)
        # Si contiene URL, no es una cita
        if "http" in citacion or "www." in citacion:
            continue
        # Si parece texto legal (largo, con artículos), sí es cita
        if any(palabra in citacion for palabra in ["artículo", "articulo", "ley", "LEY", "se establece", "dispone"]):
            hallazgos.append(Hallazgo(
                tipo="AUTO",
                severidad="WARNING",
                ubicacion=ubicacion,
                regla="comillas_citas",
                descripcion="Cita textual con comillas rectas en lugar de angulares",
                texto_original=citacion[:100],
                sugerencia="Usar comillas angulares francesas: «texto de la cita»"
            ))
    return hallazgos


# ─────────────────────────────────────────────────────────────────
# VERIFICACIÓN DE TABLAS
# ─────────────────────────────────────────────────────────────────

def verificar_tabla(tabla, num_tabla: int) -> list[Hallazgo]:
    """Verifica una tabla del documento."""
    hallazgos = []

    if len(tabla.rows) == 0:
        return hallazgos

    # Verificar que la primera fila tenga contenido en negrita
    header_row = tabla.rows[0]
    celdas_vacias = sum(1 for cell in header_row.cells if not cell.text.strip())
    if celdas_vacias > len(header_row.cells) * 0.5:
        hallazgos.append(Hallazgo(
            tipo="MANUAL",
            severidad="WARNING",
            ubicacion=f"tabla {num_tabla}",
            regla="tabla_sin_encabezado",
            descripcion="Tabla con más de la mitad de celdas vacías en la primera fila",
            texto_original="",
            sugerencia="Verificar que la primera fila sea el encabezado y tenga contenido en todas las celdas."
        ))

    # Verificar longitud del texto en celdas (si es muy largo, puede ser problema de formato)
    for row_idx, row in enumerate(tabla.rows[1:], start=2):
        for cell_idx, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if len(txt) > 500:
                hallazgos.append(Hallazgo(
                    tipo="MANUAL",
                    severidad="INFO",
                    ubicacion=f"tabla {num_tabla}, fila {row_idx}, celda {cell_idx+1}",
                    regla="tabla_celda_larga",
                    descripcion=f"Celda con {len(txt)} caracteres (puede indicar problema de formato)",
                    texto_original=txt[:100],
                    sugerencia="Verificar que el texto de la celda no esté pegado sin espacios internos."
                ))

    return hallazgos


# ─────────────────────────────────────────────────────────────────
# VERIFICACIÓN DE SECCIONES
# ─────────────────────────────────────────────────────────────────

def verificar_seccion_referencias(paragraphs, num_start: int) -> list[Hallazgo]:
    """Verifica que la sección de Referencias tenga el formato correcto."""
    hallazgos = []

    # Buscar si existe la sección
    tiene_referencias = any(
        "referencias" in p.text.lower() for p in paragraphs
    )
    if not tiene_referencias:
        hallazgos.append(Hallazgo(
            tipo="MANUAL",
            severidad="WARNING",
            ubicacion="sección Referencias",
            regla="sin_seccion_referencias",
            descripcion="El documento no tiene una sección de Referencias",
            texto_original="",
            sugerencia="Agregar una sección 'Referencias' al final del documento con las fuentes consultadas."
        ))

    return hallazgos


# ─────────────────────────────────────────────────────────────────
# ANÁLISIS COMPLETO DEL DOCUMENTO
# ─────────────────────────────────────────────────────────────────

def analizar_documento(doc_path: pathlib.Path) -> dict:
    """Analiza un documento Word y devuelve todos los hallazgos."""
    from docx import Document

    doc = Document(str(doc_path))
    todos_hallazgos = []
    siglas_definidas = set()

    # Acumular texto completo para verificación de siglas
    texto_completo = "\n".join(
        p.text for p in doc.paragraphs if p.text.strip()
    )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto_completo += "\n" + cell.text

    # Verificar siglas en todo el documento (una sola vez)
    defs, h_siglas = verificar_siglas(texto_completo)
    siglas_definidas.update(defs)
    todos_hallazgos.extend(h_siglas)

    # Verificar párrafos
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if not texto:
            continue

        ubicacion = f"párrafo {i+1}"

        # Verificar cada regla
        todos_hallazgos.extend(verificar_frases_relleno(texto, ubicacion))
        todos_hallazgos.extend(verificar_primera_persona(texto, ubicacion))
        todos_hallazgos.extend(verificar_absolutos(texto, ubicacion))
        todos_hallazgos.extend(verificar_oraciones_largas(texto, ubicacion))
        todos_hallazgos.extend(verificar_paginas(texto, ubicacion))
        todos_hallazgos.extend(verificar_articulos_bis(texto, ubicacion))
        todos_hallazgos.extend(verificar_mayusculas_total(texto, ubicacion))
        todos_hallazgos.extend(verificar_comillas_citas(texto, ubicacion))

    # Verificar tablas
    for t_idx, tabla in enumerate(doc.tables, start=1):
        todos_hallazgos.extend(verificar_tabla(tabla, t_idx))

    # Verificar sección de referencias
    todos_hallazgos.extend(verificar_seccion_referencias(doc.paragraphs, 0))

    # Clasificar por tipo
    automaticos = [h for h in todos_hallazgos if h.tipo == "AUTO"]
    manuales = [h for h in todos_hallazgos if h.tipo == "MANUAL"]

    return {
        "path": str(doc_path),
        "total": len(todos_hallazgos),
        "automaticos": automaticos,
        "manuales": manuales,
        "siglas_definidas": list(siglas_definidas),
    }


# ─────────────────────────────────────────────────────────────────
# CORRECCIÓN AUTOMÁTICA (modo --fix)
# ─────────────────────────────────────────────────────────────────

def aplicar_fixes(doc_path: pathlib.Path, output_path: pathlib.Path | None = None) -> dict:
    """Aplica las correcciones automáticas al documento."""
    from docx import Document
    from docx.shared import Pt
    from copy import deepcopy

    doc = Document(str(doc_path))
    correcciones_aplicadas = []

    if output_path is None:
        output_path = doc_path

    for para in doc.paragraphs:
        texto = para.text
        original = texto

        # 1. Corregir formato de página: 'pág. N' -> 'p. N'
        texto = PAGINA_INCORRECTO.sub(lambda m: "p. " + re.search(r"\d+", m.group()).group(), texto)

        # 2. Corregir artículos bis: 'Art. 28.5' -> 'Art. 28 Bis'
        # Primero detectar el patrón '28.5' después de un número de artículo
        def corregir_bis(m):
            art_num = m.group(1)
            return f"Art. {art_num} Bis"
        # Solo corregir si el número tiene decimales Y no es un patrón bis ya correcto
        # "28.5" -> "Art. 28 Bis", pero "28 Bis" -> sin cambio
        def fix_bis_only(m):
            antes = m.group(0)
            # Si ya dice "Bis" (correcto), no tocarlo
            if re.search(r"\bBis\b", antes, re.IGNORECASE):
                return antes
            # Es un decimal tipo "28.5" -> convertir a "28 Bis"
            m2 = re.search(r"(\d+)\.(\d+)", antes)
            if m2:
                return f"Art. {m2.group(1)} Bis"
            return antes
        texto = re.sub(r"Art\.?\s*\d+\.\d+", fix_bis_only, texto, flags=re.IGNORECASE)

        # 3. Corregir comillas rectas a angulares en citas textuales
        def angulares(m):
            citacion = m.group(1)
            if "http" in citacion or "www." in citacion:
                return m.group(0)  # No tocar URLs
            if any(p in citacion for p in ["artículo", "articulo", "ley", "LEY", "se establece", "dispone"]):
                return f"«{citacion}»"
            return m.group(0)
        texto = re.sub(r'"([^"]{20,})"', angulares, texto)

        # 4. Eliminar frases de relleno
        for pattern in FRASES_RELLENO:
            texto = pattern.sub("", texto)

        if texto != original:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = texto
            else:
                para.add_run(texto)
            correcciones_aplicadas.append({
                "antes": original[:100],
                "despues": texto[:100],
            })

    doc.save(str(output_path))

    return {
        "output": str(output_path),
        "correcciones": correcciones_aplicadas,
        "total": len(correcciones_aplicadas),
    }


# ─────────────────────────────────────────────────────────────────
# REPORTE
# ─────────────────────────────────────────────────────────────────

def generar_reporte(resultado: dict) -> str:
    """Genera un reporte de verificación en texto plano."""
    lineas = []
    lineas.append("=" * 70)
    lineas.append("REPORTE DE VERIFICACIÓN — ESTILO C-EVALUA")
    lineas.append("=" * 70)
    lineas.append(f"Documento: {resultado['path']}")
    lineas.append(f"Total hallazgos: {resultado['total']}")
    lineas.append(f"  - Automáticos (correjibles con --fix): {len(resultado['automaticos'])}")
    lineas.append(f"  - Manuales (requieren revisión humana): {len(resultado['manuales'])}")
    lineas.append(f"Siglas definidas en el documento: {', '.join(resultado['siglas_definidas']) or 'ninguna'}")
    lineas.append("")

    if resultado['automaticos']:
        lineas.append("-" * 70)
        lineas.append("HALLAZGOS AUTOMÁTICOS")
        lineas.append("-" * 70)
        agrupados = defaultdict(list)
        for h in resultado['automaticos']:
            agrupados[h.regla].append(h)
        for regla, hallazgos in agrupados.items():
            lineas.append(f"\n  [{regla}] — {len(hallazgos)} ocorrência(s)")
            for h in hallazgos[:5]:  # Mostrar máximo 5 ejemplos
                lineas.append(f"    {h.ubicacion}: {h.descripcion}")
                lineas.append(f"      → {h.sugerencia}")
            if len(hallazgos) > 5:
                lineas.append(f"    ... y {len(hallazgos)-5} más")

    if resultado['manuales']:
        lineas.append("")
        lineas.append("-" * 70)
        lineas.append("HALLAZGOS MANUALES")
        lineas.append("-" * 70)
        agrupados = defaultdict(list)
        for h in resultado['manuales']:
            agrupados[h.regla].append(h)
        for regla, hallazgos in agrupados.items():
            lineas.append(f"\n  [{regla}] — {len(hallazgos)} ocorrência(s)")
            for h in hallazgos[:5]:
                lineas.append(f"    {h.ubicacion}: {h.descripcion}")
                lineas.append(f"      → {h.sugerencia}")
            if len(hallazgos) > 5:
                lineas.append(f"    ... y {len(hallazgos)-5} más")

    if not resultado['automaticos'] and not resultado['manuales']:
        lineas.append("✅ No se encontraron hallazgos. El documento cumple con las reglas de estilo.")

    lineas.append("")
    lineas.append("=" * 70)
    lineas.append("Para aplicar correcciones automáticas:")
    lineas.append(f"  python3 {sys.argv[0]} --fix {resultado['path']}")
    lineas.append("=" * 70)

    return "\n".join(lineas)


def imprimir_hallazgos(resultado: dict):
    """Imprime los hallazgos de forma compacta en la terminal."""
    print(f"\n{'='*60}")
    print(f"  VERIFICACIÓN ESTILO C-EVALUA — {pathlib.Path(resultado['path']).name}")
    print(f"{'='*60}")
    print(f"  Total: {resultado['total']} hallazgos")
    print(f"  🔧 Auto: {len(resultado['automaticos'])}  |  👤 Manual: {len(resultado['manuales'])}")
    print(f"  Siglas definidas: {', '.join(resultado['siglas_definidas']) or 'ninguna'}")
    print()

    if resultado['automaticos']:
        print("  🔧 Automáticos:")
        agrupados = defaultdict(list)
        for h in resultado['automaticos']:
            agrupados[h.regla].append(h)
        for regla, hs in agrupados.items():
            print(f"    [{regla}] {len(hs)}x — {hs[0].descripcion[:60]}")
            for h in hs[:3]:
                print(f"      → {h.sugerencia}")

    if resultado['manuales']:
        print("  👤 Manuales:")
        agrupados = defaultdict(list)
        for h in resultado['manuales']:
            agrupados[h.regla].append(h)
        for regla, hs in agrupados.items():
            print(f"    [{regla}] {len(hs)}x — {hs[0].descripcion[:60]}")
            for h in hs[:3]:
                print(f"      → {h.sugerencia}")

    if not resultado['automaticos'] and not resultado['manuales']:
        print("  ✅ Sin hallazgos. Documento limpio.")

    print()


# ─────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Verificación y formateo de estilo C-evalua para Producto 1 FASP.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Modos:
  --check   Solo verificar y mostrar hallazgos (default si no se indica --fix)
  --fix     Verificar y aplicar correcciones automáticas
  --report  Generar reporte completo en texto plano

Ejemplos:
  python3 verificar_redaccion.py --check Producto1_Queretaro_2026.docx
  python3 verificar_redaccion.py --fix Producto1_Queretaro_2026.docx
  python3 verificar_redaccion.py --report Producto1_Queretaro_2026.docx > reporte.txt
        """
    )
    parser.add_argument("documento", help="Ruta al documento Word (.docx)")
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--check", action="store_true", help="Solo verificar (default)")
    group.add_argument("--fix", action="store_true", help="Verificar y aplicar correcciones automáticas")
    group.add_argument("--report", action="store_true", help="Generar reporte completo en texto plano")
    parser.add_argument("--output", help="Ruta de salida (para --fix)")

    args = parser.parse_args()

    doc_path = pathlib.Path(args.documento)
    if not doc_path.exists():
        print(f"ERROR: No se encontró el archivo: {doc_path}", file=sys.stderr)
        sys.exit(1)

    # Analizar
    resultado = analizar_documento(doc_path)

    if args.report:
        print(generar_reporte(resultado))
    elif args.fix:
        print(f"Aplicando correcciones a: {doc_path}")
        fix_result = aplicar_fixes(doc_path, args.output)
        print(f"\n✅ {fix_result['total']} corrección(es) aplicada(s)")
        if fix_result['correcciones']:
            print("\nCambios realizados:")
            for c in fix_result['correcciones']:
                print(f"  Antes: {c['antes'][:80]}")
                print(f"  Después: {c['despues'][:80]}")
                print()
        print(f"\nArchivo guardado en: {fix_result['output']}")
        # Volver a verificar después del fix
        print("\n--- Verificación post-fix ---")
        resultado2 = analizar_documento(pathlib.Path(fix_result['output']))
        imprimir_hallazgos(resultado2)
    else:
        # --check (default)
        imprimir_hallazgos(resultado)
        if resultado['automaticos']:
            print(f"  → Para aplicar correcciones automáticas:")
            print(f"    python3 {sys.argv[0]} --fix {doc_path}")


if __name__ == "__main__":
    main()
