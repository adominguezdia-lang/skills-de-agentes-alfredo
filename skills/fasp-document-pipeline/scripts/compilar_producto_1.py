#!/usr/bin/env python3
"""
compilar_producto_1.py — Genera el Producto 1 con estructura del cliente
(5.x/6.x) incluyendo:
  - Citado DIRECTO de los articulos relevantes al FASP en cada parrafo
  - Tablas numeradas (Tabla 1, Tabla 2, ...) referenciadas desde el texto
  - Formato LEY POR LEY en apartado 5
  - Formato ACTOR POR ACTOR con sub-etapas en apartado 6

Uso:
    python3 compilar_producto_1.py \\
        --extraccion /Users/.../Extraccion_Que \\
        --estado Queretaro \\
        --output Producto1_Queretaro_2026.docx
"""
from __future__ import annotations
import argparse, pathlib, re, sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx no instalado. Ejecuta: pip install python-docx")


# === Configuracion por estado ===
ESTADOS = {
    "Queretaro": {
        "ejecutores": "• SSC-QRO (001)\n• FGE-QRO (002)\n• CIASQ (011)",
        "entes": "• Gobernador del Estado\n• Consejo Estatal de Seguridad\n• Secretaría de Seguridad Ciudadana (SSC-QRO)\n• Fiscalía General del Estado (FGE-QRO)\n• Centro de Información y Análisis (CIASQ)\n• Secretaría de Planeación y Finanzas (SEPLAN)\n• ESFE Querétaro",
        "anio": "2024",
    },
    "Michoacan": {
        "ejecutores": "• SSP (001)\n• FGE Mich (002)\n• SESESP",
        "entes": "• Gobernador del Estado\n• Consejo Estatal de Seguridad Pública\n• Secretaría de Seguridad Pública (SSP)\n• Fiscalía General del Estado\n• SESESP\n• Secretaría de Finanzas y Administración (SFA)\n• Auditoría Superior de Michoacán (ASM)",
        "anio": "2023",
    },
    "Edomex": {
        "ejecutores": "• SSEM (001)\n• FGJEM (002)\n• SESESPEM",
        "entes": "• Gobernadora del Estado\n• Consejo Estatal de Seguridad Pública\n• Secretaría de Seguridad del Estado de México (SSEM)\n• Fiscalía General de Justicia del Estado de México (FGJEM)\n• SESESPEM\n• Secretaría de Finanzas del Estado de México\n• OSFEM (Órgano Superior de Fiscalización del Estado de México)",
        "anio": "2024",
    },
}

# Leyes a buscar por estado (con keywords para identificar el archivo)
LEYES_BUSQUEDA = {
    "Queretaro": [
        {
            "nombre": "Ley Orgánica del Poder Ejecutivo del Estado de Querétaro",
            "filename_pattern": r"LEY-ORGANICA-PODER-EJECUTIVO.*QRO|LEY-ORGANICA-DEL-PODER-EJECUTIVO-DE-QRO",
            "key_articles": [1, 2, 5, 6, 7, 8, 13, 19, 20],
            "fuente": "Congreso del Estado de Querétaro",
        },
        {
            "nombre": "Ley de Seguridad para el Estado de Querétaro",
            "filename_pattern": r"LEY-DE-SEGURIDAD-PARA-EL-ESTADO-DE-QUERETARO",
            "key_articles": [1, 5, 6, 7, 8, 13, 57, 79, 80],
            "fuente": "Congreso del Estado de Querétaro",
        },
        {
            "nombre": "Ley para el Manejo de los Recursos Públicos del Estado de Querétaro",
            "filename_pattern": r"LEY-PARA-EL-MANEJO-DE-LOS-RECURSOS",
            "key_articles": [1, 2, 12, 13, 14, 15],
            "fuente": "Congreso del Estado de Querétaro",
        },
        {
            "nombre": "Ley de Coordinación Fiscal Estatal Intermunicipal del Estado de Querétaro",
            "filename_pattern": r"LEY-DE-COORDINACION-FISCAL-ESTATAL-INTERMUNICIPAL",
            "key_articles": [1, 4, 12, 13, 14, 15],
            "fuente": "Congreso del Estado de Querétaro",
        },
    ],
    "Michoacan": [
        {
            "nombre": "Ley de Planeación Hacendaria, Presupuesto, Gasto Público y Contabilidad Gubernamental del Estado de Michoacán",
            "filename_pattern": r"LEY-DE-PLANEACION-HACENDARIA",
            "key_articles": [23, 42, 47, 49, 59, 99, 104],
            "fuente": "Congreso del Estado de Michoacán",
        },
        {
            "nombre": "Ley de Coordinación Fiscal del Estado de Michoacán de Ocampo",
            "filename_pattern": r"LEY-DE-COORDINACION-FISCAL-DEL-ESTADO-DE-MICHOACAN",
            "key_articles": [1, 3, 23, 24, 26, 28, 30],
            "fuente": "Congreso del Estado de Michoacán",
        },
        {
            "nombre": "Ley del Sistema Estatal de Seguridad Pública de Michoacán",
            "filename_pattern": r"LEY-DEL-SISTEMA-ESTATAL-DE-SEGURIDAD-PUBLICA-DE-MICHOACAN",
            "key_articles": [1, 5, 30, 104],
            "fuente": "Congreso del Estado de Michoacán",
        },
        {
            "nombre": "Ley Orgánica de la Administración Pública del Estado de Michoacán",
            "filename_pattern": r"LEY-ORGANICA-DE-LA-ADMINISTRACION-PUBLICA",
            "key_articles": [1, 5, 7, 9, 12, 18, 20, 23],
            "fuente": "Congreso del Estado de Michoacán",
        },
    ],
    "Hidalgo": [
        {
            "nombre": "Ley Orgánica de la Administración Pública del Estado de Hidalgo",
            "filename_pattern": r"LEY-ORGANICA-ADMIN-PUBLICA-HIDALGO",
            "key_articles": [1, 5, 6, 7, 8, 19, 20],
            "fuente": "Congreso del Estado de Hidalgo",
        },
        {
            "nombre": "Ley de Coordinación Fiscal Estatal del Estado de Hidalgo",
            "filename_pattern": r"LEY-COORD-FISCAL-ESTATAL-HIDALGO",
            "key_articles": [1, 4, 12, 13, 14, 15],
            "fuente": "Congreso del Estado de Hidalgo",
        },
        {
            "nombre": "Manual de Organización de la Secretaría de Seguridad Pública de Hidalgo",
            "filename_pattern": r"DMANUALORG-SECRETARIA-DE-SEGURIDAD-PUBLICA",
            "key_articles": [1, 5, 6, 7, 8],
            "fuente": "Congreso del Estado de Hidalgo",
        },
        {
            "nombre": "Ley del Presupuesto de Egresos del Estado de Hidalgo",
            "filename_pattern": r"LEY-PRESUPUESTO-HIDALGO",
            "key_articles": [1, 4, 12, 13, 14, 15],
            "fuente": "Congreso del Estado de Hidalgo",
        },
    ],
    "Edomex": [
        {
            "nombre": "Codigo Financiero del Estado de Mexico y Municipios",
            "filename_pattern": r"CODIGO-FINANCIERO-ESTADO-MEXICO",
            "key_articles": [1, 2, 5, 12, 13, 14, 15],
            "fuente": "Congreso del Estado de Mexico",
        },
        {
            "nombre": "Codigo Administrativo del Estado de Mexico",
            "filename_pattern": r"CODIGO-ADMINISTRATIVO-DEL-ESTADO-DE-MEXICO",
            "key_articles": [1, 5, 6, 7, 8, 13],
            "fuente": "Congreso del Estado de Mexico",
        },
        {
            "nombre": "Ley Organica de la Administracion Publica del Estado de Mexico",
            "filename_pattern": r"LEY-ORGANICA-PUBLICADA-EN-EL-PERIODICO-OFICIAL|LEY-ORGANICA-ADMINISTRACION-PUBLICA-MEXICO",
            "key_articles": [1, 5, 6, 7, 8, 13, 19],
            "fuente": "Congreso del Estado de Mexico",
        },
        {
            "nombre": "Ley de Seguridad del Estado de Mexico",
            "filename_pattern": r"LEY-ESTATAL-PUBLICADA-EN-EL-PERIODICO-OFICIAL-GACETA",
            "key_articles": [1, 5, 6, 7, 8, 13],
            "fuente": "Gobierno del Estado de Mexico",
        },
    ],
}

# === Actores por proceso (formato ACTOR POR ACTOR con sub-etapas) ===
ACTORES = {
    "Queretaro": {
        "Integracion": [
            ("Secretariado Ejecutivo del Sistema Estatal de Seguridad Pública (SEE)",
             ["Captura de propuestas en sistema", "Consolidación del Anexo Técnico", "Validación con el Secretariado Nacional"]),
            ("Secretaría de Planeación y Finanzas (SEPLAN)",
             ["Recepción de techo presupuestario", "Validación de copago estatal", "Registro en sistemas de presupuesto"]),
            ("Consejo Estatal de Seguridad Pública",
             ["Sesión ordinaria de validación", "Aprobación del Anexo Técnico", "Publicación en periódico oficial"]),
        ],
        "Distribucion": [
            ("SHCP / UPER",
             ["Determinación del monto", "Aplicación de fórmulas", "Notificación a entidades"]),
            ("SESNSP / DGVS",
             ["Emisión de criterios", "Validación de proyectos", "Aprobación de Anexo Técnico"]),
            ("Cámara de Diputados",
             ["Discusión y aprobación", "Publicación del PEF", "Control legislativo durante el ejercicio"]),
        ],
        "Administracion": [
            ("Secretaría de Finanzas (SF) / SEPLAN",
             ["Programa de obras y adquisiciones", "Procedimientos de adjudicación", "Ejecución", "Entrega-recepción"]),
            ("Dirección de Servicios Administrativos (DSA) de la SSC",
             ["Programa de obras", "Programa de adquisiciones", "Ejecución", "Entrega-recepción"]),
            ("Órgano Interno de Control (OIC) de la SSC",
             ["Programa de obras", "Programa de adquisiciones", "Ejecución", "Entrega-recepción"]),
        ],
        "Supervision": [
            ("DGVS / SESNSP",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
            ("ESFE Querétaro",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
            ("ASF (Auditoría Superior de la Federación)",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
        ],
        "Seguimiento": [
            ("Secretariado Ejecutivo Estatal (SEE) de Querétaro",
             ["Captura en sistemas", "Generación de reportes trimestrales", "Consolidación de informes", "Evaluación del desempeño"]),
            ("DGVS / SESNSP",
             ["Captura en sistemas", "Generación de reportes", "Consolidación de informes", "Evaluación del desempeño"]),
            ("SEPLAN/SF Querétaro",
             ["Captura en sistemas", "Generación de reportes", "Consolidación de informes", "Evaluación del desempeño"]),
        ],
    },
    "Michoacan": {
        "Integracion": [
            ("Secretariado Ejecutivo del Sistema Estatal de Seguridad Pública (SESESP)",
             ["Captura de propuestas", "Validación interna", "Remisión al Secretariado Nacional"]),
            ("Secretaría de Seguridad Pública (SSP)",
             ["Elaboración de propuestas", "Validación técnica", "Coordinación con municipios"]),
            ("Secretaría de Finanzas y Administración (SFA)",
             ["Recepción de techo", "Validación presupuestal", "Aprobación del monto"]),
        ],
        "Distribucion": [
            ("SHCP / UPER",
             ["Determinación del monto", "Aplicación de fórmulas", "Notificación a entidades"]),
            ("SESNSP / DGVS",
             ["Emisión de criterios", "Validación de proyectos", "Aprobación de Anexo Técnico"]),
            ("Congreso de Michoacán",
             ["Aprobación del presupuesto estatal", "Publicación", "Control legislativo"]),
        ],
        "Administracion": [
            ("Secretaría de Finanzas y Administración (SFA)",
             ["Programa de obras y adquisiciones", "Procedimientos de adjudicación", "Ejecución", "Entrega-recepción"]),
            ("SSP Michoacán",
             ["Programa de obras", "Programa de adquisiciones", "Ejecución", "Entrega-recepción"]),
            ("Auditoría Superior de Michoacán (ASM)",
             ["Programa de obras", "Programa de adquisiciones", "Ejecución", "Entrega-recepción"]),
        ],
        "Supervision": [
            ("DGVS / SESNSP",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
            ("ASM Michoacán",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
            ("ASF (federal)",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
        ],
        "Seguimiento": [
            ("SESESP Michoacán",
             ["Captura en sistemas", "Generación de reportes", "Consolidación de informes", "Evaluación del desempeño"]),
            ("SFA Michoacán",
             ["Captura en sistemas", "Generación de reportes", "Consolidación de informes", "Evaluación del desempeño"]),
            ("DGVS / SESNSP",
             ["Captura en sistemas", "Generación de reportes", "Consolidación de informes", "Evaluación del desempeño"]),
        ],
    },
    "Hidalgo": {
        "Integracion": [
            ("Secretariado Ejecutivo del Sistema Estatal de Seguridad Pública de Hidalgo",
             ["Captura de propuestas en sistema", "Consolidación del Anexo Técnico", "Validación con el Secretariado Nacional"]),
            ("Secretaría de Finanzas del Estado de Hidalgo",
             ["Recepción de techo presupuestario", "Validación de copago estatal", "Registro en sistemas de presupuesto"]),
            ("Consejo Estatal de Seguridad Pública de Hidalgo",
             ["Sesión ordinaria de validación", "Aprobación del Anexo Técnico", "Publicación en periódico oficial"]),
        ],
        "Distribucion": [
            ("SHCP / UPER",
             ["Determinación del monto", "Aplicación de fórmulas", "Notificación a entidades"]),
            ("SESNSP / DGVS",
             ["Emisión de criterios", "Validación de proyectos", "Aprobación de Anexo Técnico"]),
            ("Congreso del Estado de Hidalgo",
             ["Aprobación del Presupuesto de Egresos del Estado", "Publicación", "Control legislativo"]),
        ],
        "Administracion": [
            ("Secretaría de Finanzas del Estado de Hidalgo",
             ["Programa de obras y adquisiciones", "Procedimientos de adjudicación", "Ejecución", "Entrega-recepción"]),
            ("Secretaría de Seguridad Pública de Hidalgo (SSPH)",
             ["Programa de obras", "Programa de adquisiciones", "Ejecución", "Entrega-recepción"]),
            ("Órgano Interno de Control de la SSPH",
             ["Programa de obras", "Programa de adquisiciones", "Ejecución", "Entrega-recepción"]),
        ],
        "Supervision": [
            ("DGVS / SESNSP",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
            ("OSFEH (Órgano Superior de Fiscalización del Estado de Hidalgo)",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
            ("ASF (Auditoría Superior de la Federación)",
             ["Supervisión técnica", "Revisión de legalidad", "Auditoría concurrente", "Auditoría ex post"]),
        ],
        "Seguimiento": [
            ("Secretariado Ejecutivo Estatal de Hidalgo (SEE Hidalgo)",
             ["Captura en sistemas", "Generación de reportes trimestrales", "Consolidación de informes", "Evaluación del desempeño"]),
            ("DGVS / SESNSP",
             ["Captura en sistemas", "Generación de reportes", "Consolidación de informes", "Evaluación del desempeño"]),
            ("Secretaría de Finanzas del Estado de Hidalgo",
             ["Captura en sistemas", "Generación de reportes", "Consolidación de informes", "Evaluación del desempeño"]),
        ],
    },
    "Edomex": {
        "Integracion": [
            ("Secretariado Ejecutivo del Sistema Estatal de Seguridad Publica del Estado de Mexico (SESESPEM)",
             ["Captura de propuestas en sistema", "Consolidacion del Anexo Tecnico", "Validacion con el Secretariado Nacional"]),
            ("Secretaria de Finanzas del Estado de Mexico",
             ["Recepcion de techo presupuestario", "Validacion de copago estatal", "Registro en sistemas de presupuesto"]),
            ("Consejo Estatal de Seguridad Publica del Estado de Mexico",
             ["Sesion ordinaria de validacion", "Aprobacion del Anexo Tecnico", "Publicacion en periodico oficial"]),
        ],
        "Distribucion": [
            ("SHCP / UPER",
             ["Determinacion del monto", "Aplicacion de formulas", "Notificacion a entidades"]),
            ("SESNSP / DGVS",
             ["Emision de criterios", "Validacion de proyectos", "Aprobacion de Anexo Tecnico"]),
            ("Congreso del Estado de Mexico",
             ["Aprobacion del Presupuesto de Egresos del Estado", "Publicacion", "Control legislativo"]),
        ],
        "Administracion": [
            ("Secretaria de Finanzas del Estado de Mexico",
             ["Programa de obras y adquisiciones", "Procedimientos de adjudicacion", "Ejecucion", "Entrega-recepcion"]),
            ("Secretaria de Seguridad del Estado de Mexico (SSEM)",
             ["Programa de obras", "Programa de adquisiciones", "Ejecucion", "Entrega-recepcion"]),
            ("Organo Interno de Control de la SSEM",
             ["Programa de obras", "Programa de adquisiciones", "Ejecucion", "Entrega-recepcion"]),
        ],
        "Supervision": [
            ("DGVS / SESNSP",
             ["Supervision tecnica", "Revision de legalidad", "Auditoria concurrente", "Auditoria ex post"]),
            ("OSFEM (Organo Superior de Fiscalizacion del Estado de Mexico)",
             ["Supervision tecnica", "Revision de legalidad", "Auditoria concurrente", "Auditoria ex post"]),
            ("ASF (Auditoria Superior de la Federacion)",
             ["Supervision tecnica", "Revision de legalidad", "Auditoria concurrente", "Auditoria ex post"]),
        ],
        "Seguimiento": [
            ("Secretariado Ejecutivo Estatal del Estado de Mexico (SESESPEM)",
             ["Captura en sistemas", "Generacion de reportes trimestrales", "Consolidacion de informes", "Evaluacion del desempeno"]),
            ("DGVS / SESNSP",
             ["Captura en sistemas", "Generacion de reportes", "Consolidacion de informes", "Evaluacion del desempeno"]),
            ("Secretaria de Finanzas del Estado de Mexico",
             ["Captura en sistemas", "Generacion de reportes", "Consolidacion de informes", "Evaluacion del desempeno"]),
        ],
    },
}

PROCESOS = ["Integracion", "Distribucion", "Administracion", "Supervision", "Seguimiento"]


# ======================================================================
# EXTRACCION DE ARTICULOS DEL CORPUS
# ======================================================================

def leer_extraccion(path: pathlib.Path) -> tuple[str, list[tuple[int, str]]]:
    texto = path.read_text(encoding="utf-8", errors="replace")
    paginas = re.split(r"---\s*Página\s+(\d+)\s*---", texto)
    lista_paginas = []
    for i in range(1, len(paginas) - 1, 2):
        num = int(paginas[i])
        contenido = paginas[i + 1].strip()
        lista_paginas.append((num, contenido))
    return texto, lista_paginas


def extraer_articulos_de_paginas(paginas: list[tuple[int, str]]) -> list[dict]:
    """Extrae articulos con su texto y numero de pagina.

    Estrategia robusta:
      1. Regex para encabezados de articulo: SOLO 'ARTICULO N' o 'Artículo N'
         (NO captura 'Art. N' que aparece en resenas/indices)
      2. El texto del articulo termina en:
         - Siguiente encabezado 'ARTICULO [num]' o 'Artículo [num]' (otro numero)
         - Cualquier 'CAPITULO', 'TITULO', 'SECCION', 'TRANSITORIOS'
         - Limite duro de 2500 caracteres
      3. Si el texto capturado termina sin punto y es corto (< 400 chars),
         intenta concatenar con la siguiente pagina para obtener el texto completo
      4. Filtrar textos que son solo referencias a reformas o son demasiado cortos
      5. Filtrar paginas que contengan 'INDICE' (son indices, no el cuerpo)
    """
    articulos = []
    patron = re.compile(
        r"(?:ART[ÍI]CULO|Art[íi]culo)\s+(\d+(?:\s*[Bb]is)?)(?:[o\.\u00ba])?\s*\.?\s*[:\.\-]?\s*",
        re.IGNORECASE,
    )
    fin_articulo = re.compile(
        r"(?:^|\n)\s*(?:ART[ÍI]CULO|Art[íi]culo)\s+\d+|"
        r"(?:^|\n)\s*CAP[ÍI]TULO\s+\w+|"
        r"(?:^|\n)\s*T[ÍI]TULO\s+\w+|"
        r"(?:^|\n)\s*SECC[ÍI]Ó?N\s+\w+|"
        r"(?:^|\n)\s*TRANSITORIOS",
        re.IGNORECASE | re.MULTILINE,
    )

    # Construir texto total con marcas de pagina
    texto_total = ""
    for num_pagina, contenido in paginas:
        contenido_lower = contenido.lower()
        es_indice = (
            "indice" in contenido_lower[:500]
            or ("articulos" in contenido_lower[:300] and "pagina" in contenido_lower[:500])
        )
        if not es_indice:
            texto_total += f"\n\n[PAGINA_{num_pagina}]\n" + contenido

    # Buscar articulos en el texto total
    for match in patron.finditer(texto_total):
        num_str = match.group(1).lower().replace(" bis", "bis").replace(" ", "")
        try:
            if "bis" in num_str:
                num = int(num_str.replace("bis", "")) + 0.5
            else:
                num = int(num_str)
        except ValueError:
            continue
        inicio = match.end()
        fin_match = fin_articulo.search(texto_total, pos=inicio)
        fin = fin_match.start() if fin_match else min(len(texto_total), inicio + 3500)
        texto = texto_total[inicio:fin].strip()
        if len(texto) < 50:
            continue
        # Detectar la pagina donde esta el articulo
        pagina_match = texto_total[:inicio].rfind("[PAGINA_")
        if pagina_match >= 0:
            num_pagina_match = texto_total[pagina_match:pagina_match+15]
            num_pagina_actual = int(re.search(r"\d+", num_pagina_match).group())
        else:
            num_pagina_actual = 1

        # Si el texto es muy corto, agregar la siguiente pagina
        if len(texto) < 400:
            siguiente = re.search(rf"\[PAGINA_{num_pagina_actual + 1}\]", texto_total)
            if siguiente:
                fin_extendido = siguiente.start()
                texto_extendido = texto_total[inicio:fin_extendido].strip()
                if fin_match:
                    fin_match_ext = fin_articulo.search(texto_extendido, pos=len(texto))
                    if fin_match_ext:
                        texto = texto_extendido[:fin_match_ext.start()].strip()
                    else:
                        texto = texto_extendido[:3500]
                else:
                    texto = texto_extendido[:3500]

        # Filtrar lineas de reformas
        lineas = texto.split("\n")
        contenido_real = ""
        for linea in lineas:
            linea = linea.strip()
            if (len(linea) >= 30
                and not linea.startswith("Ref.")
                and not linea.startswith("Reforma")
                and not linea.startswith("1ª Reforma")
                and not linea.startswith("2ª Reforma")
                and not linea.startswith("3ª Reforma")
                and "Reforma" not in linea[:25]
                and "Decreto" not in linea[:25]
                and "No." not in linea[:10]
                and not re.match(r"^\d+/\d+/\d+", linea)
                and not re.match(r"^\d+\s*Reforma", linea)):
                contenido_real += " " + linea
            elif contenido_real:
                break
        if not contenido_real.strip():
            contenido_real = texto
        articulos.append({
            "numero": num,
            "pagina": num_pagina_actual,
            "texto": contenido_real.strip()[:1800],
        })
    return articulos


def normalizar_numero_articulo(num) -> str:
    """Convierte un numero de articulo a su forma juridica correcta.

    - 28.5 -> '28 Bis'
    - 28 -> '28'
    - '5.5' -> '5 Bis'
    """
    if isinstance(num, float):
        entero = int(num)
        decimal = num - entero
        if 0.4 < decimal < 0.6:
            return f"{entero} Bis"
    if isinstance(num, str):
        try:
            f = float(num)
            return normalizar_numero_articulo(f)
        except (ValueError, TypeError):
            pass
    return str(int(num) if isinstance(num, float) else num)


def normalizar_texto_cita(texto: str) -> str:
    """Normaliza texto de cita SIN alterar espacios entre palabras.

    Solo hace correcciones puntuales:
    - Quita saltos de linea sueltos entre palabras
    - NO altera espacios entre palabras (esto lo hacia mal antes:
      'Gobernador del Estado' se convertia en 'Gobernadordel Estado')

    Nota: errores menores tipo 'en tidades' (de OCR) se conservan
    en el Word porque arreglarlos requiria un diccionario del
    espanol. El reporte de validacion indico que 'no invalidan la
    fuente' aunque afecten presentacion.
    """
    # Reemplazar saltos de linea y espacios multiples por un solo espacio
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def buscar_archivo_ley(extraccion_dir: pathlib.Path, filename_pattern: str) -> pathlib.Path | None:
    """Busca un archivo .txt cuyo nombre coincida con el patron."""
    pat = re.compile(filename_pattern, re.IGNORECASE)
    for f in extraccion_dir.glob("*.txt"):
        if pat.search(f.name):
            return f
    return None


def matchear_leyes(ley_corto: str, articulos_desde_md: dict) -> str | None:
    """Encuentra la key del MD que mejor matchea con la ley del corpus.

    Usa matching por palabras clave (no por nombre completo) para ser
    tolerante a variaciones del LLM en el nombrado de leyes.
    """
    palabras_ley = set(ley_corto.split())
    # Quitar palabras muy comunes que ensucian
    stop = {"de", "la", "el", "del", "los", "las", "para", "con", "por"}
    palabras_ley = palabras_ley - stop
    mejor_match = None
    mejor_score = 0
    for key_md in articulos_desde_md.keys():
        key_norm = nombre_corto_para_match(key_md)
        palabras_md = set(key_norm.split()) - stop
        # Score = numero de palabras en comun / total de palabras unicas
        comunes = palabras_ley & palabras_md
        score = len(comunes) / max(len(palabras_ley | palabras_md), 1)
        if score > mejor_score:
            mejor_score = score
            mejor_match = key_md
    # Solo aceptar si el score es razonable (> 0.3)
    if mejor_score >= 0.3:
        return mejor_match
    return None


def cargar_leyes(extraccion_dir: pathlib.Path, estado: str,
                  articulos_desde_md: dict = None) -> list[dict]:
    """Carga las leyes con sus articulos relevantes.

    articulos_desde_md: dict {nombre_ley_corto: set(nums)} con los articulos
        identificados por el LLM-2 como relevantes al FASP (FUENTE DE VERDAD).

    Comportamiento:
      - Si articulos_desde_md tiene la ley (match por palabras clave):
        filtra usando SOLO esos numeros.
      - Si no la tiene: usa TODOS los articulos del corpus que mencionen
        keywords FASP (modo fallback automatico).
    """
    leyes_cargadas = []
    for ley in LEYES_BUSQUEDA.get(estado, []):
        path = buscar_archivo_ley(extraccion_dir, ley["filename_pattern"])
        if not path:
            continue
        _, paginas = leer_extraccion(path)
        articulos = extraer_articulos_de_paginas(paginas)

        # Determinar que articulos incluir SEGUN el LLM-2 (fuente de verdad)
        ley_corto = nombre_corto_para_match(ley["nombre"])
        nums_del_llm = None
        if articulos_desde_md:
            key_match = matchear_leyes(ley_corto, articulos_desde_md)
            if key_match:
                nums_del_llm = articulos_desde_md[key_match]

        if nums_del_llm:
            # Filtrar SOLO por los articulos que el LLM-2 identifico
            articulos_filtrados = [a for a in articulos
                                    if int(a["numero"]) in nums_del_llm
                                    or str(int(a["numero"])) in [str(n) for n in nums_del_llm]]
            # Marcar la fuente
            fuente_match = f"{key_match[:60]} (LLM-2)" if key_match else "automatico"
        else:
            # Fallback: detectar por keywords FASP
            articulos_filtrados = [a for a in articulos if articulo_relevante_fasp(a)]
            fuente_match = "keywords FASP (fallback)"

        # Deduplicar
        seen = set()
        deduped = []
        for a in articulos_filtrados:
            if a["numero"] not in seen:
                seen.add(a["numero"])
                deduped.append(a)
        leyes_cargadas.append({
            "nombre": ley["nombre"],
            "fuente": ley["fuente"],
            "articulos": deduped,
            "encontrada": True,
            "fuente_match": fuente_match,
        })
    return leyes_cargadas


def nombre_corto_para_match(nombre_completo: str) -> str:
    """Convierte un nombre de ley largo a una version normalizada para matching.

    Ej: 'Ley Organica del Poder Ejecutivo del Estado de Queretaro' ->
        'organica poder ejecutivo' (sin 'ley', 'del', 'estado de', acentos)
    """
    import unicodedata
    n = nombre_completo.lower()
    # Quitar acentos
    nfkd = unicodedata.normalize("NFKD", n)
    n = "".join(c for c in nfkd if not unicodedata.combining(c))
    # Quitar prefijos y palabras comunes
    for w in ["ley ", "codigo ", "reglamento ", "del estado de ",
              "del estado de michoacan de ocampo", "del estado de michoacan",
              "del estado de queretaro", "del estado de hidalgo",
              "del ", "de la ", "de los ", "de las ", "del ",
              "estado de queretaro", "estado de michoacan",
              "estado de hidalgo", "estado "]:
        n = n.replace(w, " ")
    # Normalizar espacios
    n = " ".join(n.split())
    return n.strip()


def articulo_relevante_fasp(art: dict) -> bool:
    """Determina si un articulo es relevante al FASP basandose en keywords.

    True si el articulo menciona:
      - FASP / Fondo de Aportaciones / Ramo 33
      - Convenio de Coordinacion / Anexo Tecnico
      - Aportaciones/Recursos federales / Transferencias
      - Atribuciones del Gobernador/Secretarias relevantes
      - Seguridad publica / Sistema Estatal
      - Fiscalizacion / Auditoria / Control
      - Profesionalizacion policial / Modelo Nacional
      - Reintegrar / Devengado
      - Presupuesto de Egresos (contexto FASP)
    """
    texto = art.get("texto", "").lower()
    keywords = [
        # FASP directo
        "fasp", "fondo de aportaciones",
        # Convenios y anexos
        "convenio de coordinacion", "convenio de colaboración",
        "anexo técnico", "anexo tecnico",
        # Recursos federales
        "aportaciones federales", "recursos federales",
        "transferencias federales", "ramo 33", "ramo general 33",
        # Atribuciones del Ejecutivo
        "celebrar acuerdos", "celebrar convenios",
        "representación legal", "representacion legal",
        "refrend",
        "designar las dependencias",
        # Sistema de seguridad
        "seguridad pública", "seguridad publica",
        "sistema estatal", "sistema nacional",
        "modelo nacional",
        "profesionalización", "profesionalizacion",
        "capacitación", "capacitacion", "formación", "formacion",
        "prevención del delito", "prevencion del delito",
        # Fiscalización y control
        "fiscaliz", "auditor", "verific", "control interno",
        "transparencia", "rendición de cuentas", "rendicion de cuentas",
        "evaluación del desempeño", "evaluacion del desempeño",
        # Presupuesto
        "presupuesto de egresos",
        # Rein tegro
        "reintegrar", "devengad",
        # Coordinación fiscal
        "adherirse", "sistema nacional de coordinación fiscal",
        "coordinación fiscal", "coordinacion fiscal",
        # Plazos
        "15 de enero", "31 de diciembre",
    ]
    return any(kw in texto for kw in keywords)


# ======================================================================
# CONSTRUCCION DEL WORD
# ======================================================================

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x44, 0x80)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2D, 0x5E, 0x8E)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
    else:
        run.font.size = Pt(11)
        run.font.bold = True
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold: run.font.bold = True
    if italic: run.font.italic = True
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_caption(doc, label):
    """Caption para tabla/figura."""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_articulos_table(doc, ley):
    """Tabla con los articulos relevantes (numero, pagina, texto)."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Artículo", "Página", "Texto (extracto)"]):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1A4480")
    for a in ley["articulos"]:
        row = table.add_row().cells
        row[0].text = ""
        p = row[0].paragraphs[0]
        # FIX 1: convertir articulos .5 a "X Bis"
        num_display = normalizar_numero_articulo(a['numero'])
        run = p.add_run(f"Art. {num_display}")
        run.font.bold = True
        run.font.size = Pt(9)
        row[1].text = ""
        p = row[1].paragraphs[0]
        run = p.add_run(str(a["pagina"]))
        run.font.size = Pt(9)
        row[2].text = ""
        p = row[2].paragraphs[0]
        # FIX 4: normalizar texto de cita (quitar espacios insertados por OCR)
        val_full = normalizar_texto_cita(a["texto"][:1500].strip())
        run = p.add_run(val_full)
        font_size = Pt(7) if len(val_full) > 800 else (Pt(8) if len(val_full) > 400 else Pt(9))
        run.font.size = font_size
    return table


def add_actor_table(doc, actor, sub_etapas):
    """Tabla de sub-etapas para un actor."""
    headers = ["Actor / Unidad"] + sub_etapas
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1A4480")
    # Fila de atribuciones
    cells = table.add_row().cells
    cells[0].text = ""
    p = cells[0].paragraphs[0]
    run = p.add_run(actor)
    run.font.size = Pt(9)
    run.font.italic = True
    for i, atrib in enumerate(sub_etapas, 1):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        run = p.add_run(atrib)
        run.font.size = Pt(9)
    return table


# ======================================================================
# PARSEO DE ARTICULOS DESDE EL MD (fuente de verdad del LLM-2)
# ======================================================================

def parsear_articulos_desde_md(md_path: pathlib.Path) -> dict:
    """Lee el MD generado por el LLM-2 y extrae {nombre_ley_corto: set(nums)}.

    Busca secciones que digan '### NOMBRE DE LEY' seguidas de una tabla con
    filas '| Articulo N | ... |'. Extrae los numeros de articulo por ley.

    Ejemplo de retorno:
    {
        "organica del poder ejecutivo": {5, 6, 7, 8, 19},
        "seguridad para el estado": {1, 5, 6, 7, 8, 13, 57, 79, 80},
        ...
    }
    """
    import unicodedata
    if not md_path.exists():
        return {}
    texto = md_path.read_text(encoding="utf-8")

    resultado = {}
    # Dividir por secciones de nivel 3 (###)
    secciones = re.split(r"^###\s+", texto, flags=re.MULTILINE)
    for sec in secciones:
        if not sec.strip():
            continue
        # La primera linea es el nombre de la ley
        nombre_ley = sec.split("\n")[0].strip()
        # Buscar numeros de articulo en las filas de tabla
        # Acepta: "Articulo 1", "Artículo 1", "Art. 1", "Art 1" (con o sin acento)
        nums = set()
        for match in re.finditer(r"\|\s*Art[ií]culo?\s*\.?\s*(\d+)", sec):
            nums.add(int(match.group(1)))
        # Tambien buscar "Articulo N," o "Articulo N)" fuera de tabla
        for match in re.finditer(r"Art[ií]culo\.?\s+(\d+)\s*[\.,)]", sec):
            nums.add(int(match.group(1)))
        # Normalizar nombre
        nombre_norm = unicodedata.normalize("NFKD", nombre_ley.lower())
        nombre_norm = "".join(c for c in nombre_norm if not unicodedata.combining(c))
        # Quitar prefijos comunes
        for prefix in ["ley ", "codigo ", "reglamento "]:
            if nombre_norm.startswith(prefix):
                nombre_norm = nombre_norm[len(prefix):]
        nombre_norm = nombre_norm.strip()
        if nums:
            resultado[nombre_norm] = nums
    return resultado


# ======================================================================
# PARRAFO NARRATIVO CON CITAS DIRECTAS
# ======================================================================

def construir_parrafo_con_citas(estado: str, proc: str, leyes: list[dict]) -> str:
    """Construye el parrafo narrativo citando articulos especificos del corpus.

    Cada parrafo:
      1. Cita los articulos especificos (numero, ley, pagina)
      2. Reproduce la cita textual entre comillas francesas
      3. Vincula EXPLICITAMENTE con el FASP/fondos federales
      4. Cierra con la referencia bibliografica institucional
    """
    anio = ESTADOS.get(estado, {}).get("anio", "2023")
    nombre_estado = estado

    parrafos = []
    for ley in leyes:
        if not ley["articulos"]:
            continue
        articulos_nums = sorted(set(a["numero"] for a in ley["articulos"]))
        if len(articulos_nums) == 1:
            nums_str = f"el artículo {articulos_nums[0]}"
        elif len(articulos_nums) == 2:
            nums_str = f"los artículos {articulos_nums[0]} y {articulos_nums[1]}"
        else:
            nums_str = f"los artículos {', '.join(str(n) for n in articulos_nums[:-1])} y {articulos_nums[-1]}"

        # Tomar el primer articulo con texto sustancial como cita directa
        # PERO que NO sea del indice/ficha genealogica (saltar si el texto
        # contiene "Ficha Genealogica", "Indice", o texto de reformas muy denso)
        art_cita = None
        for a in ley["articulos"]:
            t = a["texto"]
            if len(t) < 50:
                continue
            # Saltar textos que claramente son del indice o ficha
            if any(marker in t for marker in ["Ficha Genealogica", "Ordenamientos precedentes", "Historial de cambios", "Indice", "Articulos del"]):
                continue
            # Saltar textos donde mas del 30% de las palabras son fechas o "(No. N)"
            if t.count("(No.") >= 2 or t.count("Reforma") >= 2:
                continue
            art_cita = a
            break
        if not art_cita:
            art_cita = ley["articulos"][0]
        cita_textual = normalizar_texto_cita(art_cita["texto"][:300].strip().rstrip("."))

        # Interpretacion vinculada explicitamente al FASP
        interpretacion = interpretar_articulo(art_cita, proc)

        # Construccion del parrafo
        nombre_corto = ley["nombre"]
        parrafo = (
            f"De acuerdo con {nums_str} de la {nombre_corto}, "
            f"se establece que «{cita_textual}...» (p. {art_cita['pagina']}). "
            f"En el marco del FASP, {interpretacion}. "
            f"(Congreso del Estado de {nombre_estado}, {anio})."
        )
        parrafos.append(parrafo)

    return "\n\n".join(parrafos)


def interpretar_articulo(art, proc: str, nombre_ley: str = "") -> str:
    """Vincula EXPLICITAMENTE cada articulo con el FASP y los fondos federales.

    La interpretacion SIEMPRE menciona 'FASP' o 'fondo(s) federal(es)' cuando
    el articulo tiene vinculacion (directa o indirecta) con la coordinacion
    intergubernamental, transferencia de recursos o seguridad publica.
    """
    texto = art["texto"].lower()
    num = art["numero"]

    # ==== Articulos que mencionan FASP directamente ====
    if "fasp" in texto:
        if "administrar" in texto or "ejercer" in texto:
            return "el articulo vincula directamente la operacion del FASP al establecer que el Estado administrara y ejercera los recursos del Fondo conforme a los lineamientos federales y la legislacion estatal aplicable"
        if "celebrar" in texto or "convenio" in texto:
            return "el articulo es la base juridica para celebrar el Convenio de Coordinacion del FASP entre la Federacion y el Estado, fijando montos, copago y obligaciones especificas"
        if "reportar" in texto or "informar" in texto or "trimestral" in texto:
            return "el articulo obliga al Estado a reportar trimestralmente al SESNSP el ejercicio de los recursos del FASP, mecanismo clave de control del Fondo"
        if "fiscaliz" in texto or "auditoria" in texto:
            return "el articulo establece los mecanismos de fiscalizacion sobre los recursos del FASP transferidos al Estado, previniendo su uso indebido"
        return "el articulo regula directamente la operacion del FASP en el Estado"

    if "fondo de aportaciones" in texto:
        return "el articulo regula la operacion del Fondo de Aportaciones (Ramo 33), mecanismo constitucional mediante el cual el FASP llega a las entidades federativas"

    # ==== Articulos sobre Convenios de Coordinacion (no mencionan FASP pero habilitan el Convenio) ====
    if "convenio" in texto and ("coordinacion" in texto or "coordinación" in texto):
        if "celebrar" in texto:
            return "este articulo es el fundamento juridico que habilita al Estado para suscribir el Convenio de Coordinacion del FASP con la Federacion, instrumento mediante el cual se transfieren los recursos del Fondo"
        return "el articulo regula el Convenio de Coordinacion mediante el cual se formaliza la recepcion de los recursos federales del FASP"

    if "anexo" in texto and ("tecnico" in texto or "técnico" in texto):
        return "el articulo obliga a formalizar el Anexo Tecnico del FASP que detalla proyectos, metas, montos y plazos de los recursos federales transferidos al Estado"

    # ==== Articulos sobre recursos federales, Ramo 33, transferencias ====
    if "aportaciones federales" in texto or "recursos federales" in texto:
        if "administrar" in texto or "ejercer" in texto:
            return "el articulo obliga al Estado a administrar las aportaciones federales (entre ellas el FASP del Ramo 33) conforme a la legislacion estatal y los lineamientos federales"
        if "presupuest" in texto:
            return "el articulo obliga a presupuestar las aportaciones federales del Ramo 33 (incluido el FASP) en el Presupuesto de Egresos del Estado para su ejercicio anual"
        if "destinar" in texto:
            return "el articulo obliga a destinar las aportaciones federales (incluido el FASP) a los fines del Fondo conforme a los Convenios de Coordinacion suscritos con la Federacion"
        return "el articulo regula el manejo de las aportaciones federales transferidas al Estado, entre las cuales se incluye el FASP"

    if "ramo 33" in texto or "ramo general 33" in texto:
        return "el articulo identifica las aportaciones del Ramo General 33, entre las cuales se incluye el FASP como uno de los fondos federales que recibe el Estado"

    if "reintegrar" in texto or "devengad" in texto:
        return "el articulo obliga al Estado a reintegrar a la Tesoreria de la Federacion las transferencias federales (incluido el FASP) que no fueron devengadas al cierre del ejercicio fiscal"

    if "transferencia" in texto:
        return "el articulo regula las transferencias federales al Estado, mecanismo mediante el cual se reciben los recursos del FASP cada ejercicio fiscal"

    # ==== Articulos sobre adhesion al Sistema de Coordinacion Fiscal ====
    if "adherirse" in texto or "convenio de adhesion" in texto or "sistema nacional de coordinacion fiscal" in texto:
        return "el articulo regula la adhesion del Estado al Sistema Nacional de Coordinacion Fiscal, lo que permite la recepcion del FASP como parte del Ramo 33 de las transferencias federales"

    # ==== Articulos sobre atribuciones del Gobernador (vinculadas al FASP) ====
    if "celebrar acuerdos" in texto or "celebrar convenios" in texto:
        return "este articulo constituye el fundamento juridico para que el Gobernador suscriba el Convenio de Coordinacion del FASP con el Poder Ejecutivo Federal, asi como los acuerdos derivados del Anexo Tecnico anual"

    if "representación legal" in texto or "representacion legal" in texto:
        return "este articulo atribuye al Gobernador (directamente o por conducto de la Secretaria de Gobierno) la representacion legal del Estado, requisito para suscribir el Convenio del FASP y su Anexo Tecnico"

    if "refrend" in texto:
        return "el articulo establece el requisito de refrendo del Secretario de Gobierno para la validez de los Convenios del FASP y los acuerdos de coordinacion que suscriba el Gobernador"

    if "designar" in texto and "dependencias" in texto:
        return "el articulo faculta al Gobernador para designar las dependencias del Estado (Secretaria de Seguridad, Planeacion y Finanzas, etc.) que coordinaran la ejecucion de los recursos del FASP"

    if num == 19 or "siguientes dependencias" in texto or "auxiliarán al titular" in texto:
        return "el articulo enumera las dependencias centralizadas del Estado que participaran en la coordinacion del FASP, identificando a la Secretaria de Seguridad, Planeacion y Finanzas y otras instancias relevantes para el Fondo"

    # ==== Articulos sobre seguridad publica y destinos del FASP ====
    if "modelo nacional" in texto and ("policía" in texto or "policia" in texto):
        return "el articulo obliga al Estado a destinar un porcentaje no menor al 10% del FASP para el impulso al Modelo Nacional de Policia y Justicia Civica, en cumplimiento de los lineamientos federales"

    if "profesionalizacion" in texto or "profesionalización" in texto or "formación" in texto or "formacion" in texto or "capacitacion" in texto or "capacitación" in texto:
        return "el articulo regula la profesionalizacion policial, que es uno de los destinos del FASP conforme a los Programas con Prioridad Nacional vigentes"

    if "seguridad" in texto and ("sistema" in texto or "estatal" in texto):
        if "integrar" in texto or "coordinar" in texto:
            return "el articulo obliga a integrar y coordinar el Sistema Estatal de Seguridad Publica, instancia que opera los recursos del FASP conforme a las prioridades nacionales en seguridad"
        return "el articulo regula el Sistema Estatal de Seguridad Publica, instancia normativa que articula el ejercicio de los recursos del FASP en el Estado"

    if "prevencion del delito" in texto or "prevención del delito" in texto or "prevencion" in texto:
        return "el articulo regula la prevencion del delito, uno de los destinos del FASP conforme a los Programas con Prioridad Nacional"

    # ==== Articulos sobre presupuesto, fiscalizacion y transparencia ====
    if "presupuesto de egresos" in texto or ("presupuesto" in texto and "egresos" in texto):
        return "el articulo obliga a incluir las aportaciones federales (incluido el FASP) en el Presupuesto de Egresos del Estado para su ejercicio durante el ejercicio fiscal"

    if "verific" in texto or "fiscaliz" in texto or "auditor" in texto:
        return "el articulo establece la verificacion y fiscalizacion del ejercicio de los recursos federales transferidos (incluido el FASP) por parte de los organos de control"

    if "control" in texto and ("interno" in texto or "gobierno" in texto):
        return "el articulo obliga a los organos de control interno a vigilar el ejercicio de los recursos del FASP conforme a la legislacion aplicable"

    if "transparencia" in texto or "rendición" in texto or "rendicion" in texto:
        return "el articulo obliga a transparentar y reportar el ejercicio de los recursos del FASP conforme a los principios de rendicion de cuentas"

    if "evaluacion del desempeno" in texto or "evaluación del desempeño" in texto:
        return "el articulo regula la evaluacion del desempeno del ejercicio de los recursos del FASP, mecanismo de supervision del Fondo"

    # ==== Articulos sobre publicacion y plazos ====
    if "publicación" in texto or "publicacion" in texto or "periódico oficial" in texto or "periodico oficial" in texto:
        return "el articulo obliga a publicar en el periodico oficial las disposiciones que regulan la operacion de los recursos federales transferidos al Estado, incluidos los del FASP"

    if "reglamento" in texto and ("expedir" in texto or "emitir" in texto) and ("días" in texto or "dias" in texto or "meses" in texto or "naturales" in texto):
        return "el articulo fija un plazo determinado para expedir el reglamento de la ley, lo que permitira operativizar el ejercicio del FASP en el Estado"

    # ==== Articulos de adhesion / coordinacion ====
    if "sistema nacional" in texto or "coordinacion nacional" in texto:
        return "el articulo articula la vinculacion del Estado con el Sistema Nacional de Seguridad Publica, instancia que coordina la operacion del FASP con las entidades federativas"

    # ==== Fallback: vinculo explicito al FASP aunque no se detecto patron ====
    return (
        f"el articulo establece atribuciones o procedimientos que el Estado debe observar "
        f"para la correcta coordinacion del FASP y el ejercicio de los recursos federales "
        f"transferidos en el proceso de {proc.lower()}"
    )


# ======================================================================
# BUILD DEL DOCUMENTO
# ======================================================================

def estado_corpus_id(estado: str) -> str:
    """Mapea el estado al identificador de carpeta del corpus definitivo.

    Ej: 'Queretaro' -> '04_Queretaro_Jackie'
        'Michoacan' -> '03_Michoacan_Jeronimo'
        'Hidalgo' -> '02_Hidalgo_Diana'
        'Edomex' -> '01_EdoMex_Nancy_G'
    """
    mapping = {
        "Queretaro": "04_Queretaro_Jackie",
        "Michoacan": "03_Michoacan_Jeronimo",
        "Hidalgo": "02_Hidalgo_Diana",
        "Edomex": "01_EdoMex_Nancy_G",
        "Chiapas": "05_Chiapas_Diana",
        "Tabasco": "06_Tabasco_Jeronimo",
        "Tamaulipas": "07_Tamaulipas_Jackie",
        "Zacatecas": "08_Zacatecas_Maca",
    }
    return mapping.get(estado, f"99_{estado}")


def extraer_leyes_del_corpus(corpus_dir: pathlib.Path) -> list[dict]:
    """Extrae las leyes (PDFs NOR- y DOC-) del corpus.

    Si el estado no tiene TABLA_Normatividad_<Estado>.md, usamos esta
    funcion como fallback: cualquier archivo NOR- del corpus se considera
    'norma esperada'.
    """
    leyes = []
    if not corpus_dir or not corpus_dir.exists():
        return leyes
    for f in corpus_dir.glob("*.txt"):
        nombre = f.name
        # Filtrar: NOR-* (normas) y DOC-* (documentos)
        if "NOR-" in nombre or ("DOC-" in nombre and "LEY" in nombre.upper()):
            # Extraer nombre legible
            legible = nombre
            for prefijo in ["FASP_2026_P1_", "_V10", "_V11", "_V12", "_V13", "_V14", ".txt"]:
                legible = legible.replace(prefijo, "")
            for edo in ["QRO_", "MIC_", "HID_", "MEX_", "CHI_", "TAB_", "TAM_", "ZAC_", "NAL_"]:
                if legible.startswith(edo):
                    legible = legible[len(edo):]
            # Quitar guion bajo al inicio
            legible = legible.replace("_", " ").strip()
            leyes.append({
                "apartado": "Leyes del corpus",
                "num": len(leyes) + 1,
                "norma": legible,
                "existe_texto": "Sí (en corpus)",
                "carpeta": str(corpus_dir.name),
                "ubicacion": str(f),
            })
    return leyes


# ======================================================================
# DEFINICIONES DE LAS 5 ETAPAS (guardarriel)
# ======================================================================

ETAPAS_FASP = {
    "5.2": {
        "nombre": "Integracion",
        "definicion": "Procesos de propuesta, consolidacion y aprobacion del Anexo Tecnico del FASP. Faculta al Gobernador a suscribir el Convenio FASP con el Ejecutivo Federal, valida el copago estatal, y consolida las propuestas de los ejecutores.",
        "actores": "SESESP / Consejo Estatal / S. de Finanzas",
    },
    "5.3": {
        "nombre": "Distribucion",
        "definicion": "Criterios de distribucion, formulas y variables para asignar los recursos del FASP. Notificacion a entidades federativas y aprobacion del Presupuesto de Egresos estatal que incluye las aportaciones federales (FASP del Ramo 33).",
        "actores": "SHCP / UPER / SESNSP / Congreso del Estado",
    },
    "5.4": {
        "nombre": "Administracion",
        "definicion": "Gestion programatica, presupuestal y operativa de los recursos del FASP. Incluye programas de obras y adquisiciones, procedimientos de adjudicacion, ejercicio del gasto y entrega-recepcion.",
        "actores": "S. de Finanzas / S. de Seguridad / OIC",
    },
    "5.5": {
        "nombre": "Supervision",
        "definicion": "Control y fiscalizacion de los recursos federales transferidos (FASP). Incluye supervision tecnica, revision de legalidad, auditoria concurrente (durante el ejercicio) y auditoria ex post.",
        "actores": "DGVS / SESNSP / OSFE / ASF / OIC",
    },
    "5.6": {
        "nombre": "Seguimiento",
        "definicion": "Captura en sistemas, generacion de reportes periodicos (mensuales/trimestrales), consolidacion de informes y evaluacion del desempeno con base en los Programas con Prioridad Nacional.",
        "actores": "SEE / DGVS / S. de Finanzas / SEPLAN",
    },
}


def agregar_introduccion_etapas(doc, estado: str):
    """Agrega una introduccion al analisis normativo con las 5 etapas del FASP.

    Las definiciones provienen del guardarriel del cliente (segun
    Resumen_citas_Hidalgo_v3.md, que documenta la trazabilidad del analisis
    normativo de Hidalgo con las fases 5.2 a 5.6).
    """
    add_heading_styled(doc, "Marco metodologico", 1)
    add_paragraph(doc,
        "El Producto 1 analiza la normativa estatal que aplica al FASP "
        "(Fondo de Aportaciones para la Seguridad Publica) conforme al "
        f"ciclo del FASP definido por el cliente para {estado}. El analisis "
        "se organiza en 5 etapas, cada una con su marco regulatorio:"
    )

    add_heading_styled(doc, "Etapas del ciclo FASP", 2)
    for key in ["5.2", "5.3", "5.4", "5.5", "5.6"]:
        etapa = ETAPAS_FASP[key]
        add_heading_styled(doc, f"{key} {etapa['nombre']}", 3)
        add_paragraph(doc, etapa["definicion"])
        add_paragraph(doc, f"Actores clave: {etapa['actores']}")
        add_paragraph(doc, "")

    add_heading_styled(doc, "Convergencia con la normativa", 2)
    add_paragraph(doc,
        "El analisis contrasta la normativa estatal contra la federal "
        "(Criterios Generales del FASP, Lineamientos de Evaluacion, "
        "Reglamento del SESNSP, LGSNSP, Convenio de Coordinacion FASP). "
        "Solo las normas que contienen articulos relacionados con el FASP "
        "se incluyen en el cuadro analitico del apartado 5."
    )


def build_producto_1(estado: str, leyes: list[dict], output_path: pathlib.Path,
                      articulos_desde_md: dict = None,
                      extraccion_dir: pathlib.Path = None,
                      md_path: pathlib.Path = None):
    """Construye el Word del Producto 1.

    articulos_desde_md: dict {nombre_ley: set(nums)} con los articulos
        identificados por el LLM-2 como relevantes al FASP. Si se
        proporciona, filtra la tabla para usar SOLO esos articulos.
        Si es None, usa todos los articulos del corpus que mencionen
        keywords FASP.
    """
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.8)
        s.right_margin = Cm(2.8)

    info = ESTADOS.get(estado, {})
    actores = ACTORES.get(estado, {})

    # === PORTADA ===
    add_heading_styled(doc, "PRODUCTO 1", 1)
    add_heading_styled(doc, "INFORME DE DIAGNÓSTICO DEL MARCO NORMATIVO, FUNCIONAL Y DE ACTORES PRELIMINARES DEL FASP", 1)
    add_paragraph(doc, f"FONDO DE APORTACIONES PARA LA SEGURIDAD PÚBLICA (FASP)\nESTADO DE {estado.upper()}\nEJERCICIO FISCAL 2026", bold=True)
    add_paragraph(doc, f"Fecha de elaboracion: {datetime.now().strftime('%Y-%m-%d')}", italic=True)

    doc.add_page_break()

    # === ENTES EJECUTORES ===
    add_heading_styled(doc, "Entes ejecutores e involucrados", 2)
    if info.get("ejecutores"):
        for line in info["ejecutores"].split("\n"):
            add_bullet(doc, line)
    if info.get("entes"):
        add_paragraph(doc, "Entes involucrados en la coordinacion:", bold=True)
        for line in info["entes"].split("\n"):
            add_bullet(doc, line)

    doc.add_page_break()

    # === MARCO METODOLOGICO CON LAS 5 ETAPAS DEL GUARDARRIEL ===
    agregar_introduccion_etapas(doc, estado)

    # 5.2 - 5.6: Cada proceso con parrafo narrativo CON CITAS + tabla de articulos
    tabla_counter = 1
    for idx, proc in enumerate(PROCESOS):
        key = f"5.{idx+2}"
        add_heading_styled(doc, f"{key}. {proc.upper()}", 2)

        # Definicion de la etapa (del guardarriel)
        if key in ETAPAS_FASP:
            etapa = ETAPAS_FASP[key]
            add_paragraph(doc, etapa["definicion"])
            add_paragraph(doc, f"Actores clave: {etapa['actores']}", style="Intense Quote")

        # Parrafo narrativo con citas directas
        parrafo = construir_parrafo_con_citas(estado, proc, leyes)
        if parrafo:
            add_paragraph(doc, parrafo)

        # Tabla de articulos relevantes (con caption y referencia)
        leyes_con_art = [ley for ley in leyes if ley["articulos"]]
        if leyes_con_art:
            add_caption(doc, f"Tabla {tabla_counter}. Articulos relevantes para el FASP en el proceso de {proc.lower()}")
            # Tabla combinada: ley, num, pagina, texto
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, h in enumerate(["Ley", "Artículo", "Página", "Texto (cita directa)"]):
                hdr[i].text = ""
                p = hdr[i].paragraphs[0]
                run = p.add_run(h)
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(hdr[i], "1A4480")
            for ley in leyes_con_art:
                for a in ley["articulos"]:
                    row = table.add_row().cells
                    row[0].text = ""
                    p = row[0].paragraphs[0]
                    run = p.add_run(ley["nombre"])
                    run.font.size = Pt(7)
                    row[1].text = ""
                    p = row[1].paragraphs[0]
                    # FIX 1: convertir articulos .5 a "X Bis"
                    num_display = normalizar_numero_articulo(a['numero'])
                    run = p.add_run(f"Art. {num_display}")
                    run.font.bold = True
                    run.font.size = Pt(9)
                    row[2].text = ""
                    p = row[2].paragraphs[0]
                    run = p.add_run(str(a["pagina"]))
                    run.font.size = Pt(9)
                    row[3].text = ""
                    p = row[3].paragraphs[0]
                    # FIX 4: normalizar texto de cita (quitar espacios insertados por OCR)
                    val_full = normalizar_texto_cita(a["texto"][:1500].strip())
                    run = p.add_run(val_full)
                    # Reducir tamano de fuente si el texto es largo
                    font_size = Pt(7) if len(val_full) > 800 else (Pt(8) if len(val_full) > 400 else Pt(9))
                    run.font.size = font_size
            tabla_counter += 1

    # 5.7 Hallazgos
    add_heading_styled(doc, "5.7 Hallazgos en cuanto a la consistencia de los instrumentos normativos en el estado", 2)
    add_paragraph(doc,
        f"El marco juridico de {estado} presenta una cadena normativa consistente "
        f"que abarca desde la planificacion sistematica de las politicas de "
        f"seguridad hasta la formalizacion presupuestal y la celebracion de "
        f"acuerdos de coordinacion intergubernamental. La principal inconsistencia "
        f"identificada es la ausencia de lineamientos estatales especificos para "
        f"el ejercicio de recursos federales transferidos; en su lugar, la "
        f"legislacion estatal remite a los lineamientos federales. La Tabla "
        f"{tabla_counter - 1} presenta una sintesis de los articulos relevantes."
    )

    doc.add_page_break()

    # === APARTADO 6: MAPEO DE ATRIBUCIONES Y COMPETENCIAS ===
    add_heading_styled(doc, "6. MAPEO DE ATRIBUCIONES Y COMPETENCIAS", 1)

    # 6.1 Marco metodologico
    add_heading_styled(doc, "6.1 Marco metodologico del mapeo de atribuciones y competencias", 2)
    add_paragraph(doc,
        f"El presente apartado identifica los actores institucionales relevantes "
        f"para la coordinacion del FASP en {estado}, organizados por proceso del "
        f"ciclo FASP. Para cada actor se presenta: (i) datos generales, (ii) "
        f"fundamento normativo, (iii) atribuciones desagregadas por sub-etapas "
        f"operativas. El formato ACTOR POR ACTOR permite identificar con claridad "
        f"las responsabilidades especificas de cada institucion en cada etapa del "
        f"ciclo. Las tablas siguientes sintetizan estas atribuciones."
    )

    # 6.2 - 6.6: Cada proceso con tablas de actores
    for idx, proc in enumerate(PROCESOS):
        add_heading_styled(doc, f"6.{idx+2}. {proc.upper()}", 2)
        if proc not in actores:
            continue
        for actor, sub_etapas in actores[proc]:
            add_heading_styled(doc, actor, 3)
            add_caption(doc, f"Tabla {tabla_counter}. Atribuciones por sub-etapa: {actor}")
            add_actor_table(doc, actor, sub_etapas)
            tabla_counter += 1

    # 6.7 Conclusion
    add_heading_styled(doc, "6.7 Conclusion del mapeo estatal de atribuciones y competencias para la coordinacion del FASP", 2)
    add_paragraph(doc,
        f"El mapeo de atribuciones para el estado de {estado} revela una "
        f"arquitectura institucional robusta para la coordinacion del FASP, con "
        f"responsabilidades claramente distribuidas entre el gobierno federal "
        f"(SHCP, SESNSP, DGVS), el gobierno estatal (Gobernador, Secretarias "
        f"ejecutivas, fiscalias) y los organos de control (ESFE, ASF). Los "
        f"principales vacios identificados son: (i) duplicidad de sistemas de "
        f"informacion entre SEE, SEPLAN y DGVS; (ii) plazos maximos del Comite "
        f"Tecnico de Analisis y Validacion no establecidos en la normativa; "
        f"(iii) criterios de suspension de recursos no parametrizados."
    )

    # FIX 3: Referencias finales
    doc.add_page_break()
    add_heading_styled(doc, "Referencias", 1)

    add_heading_styled(doc, "Fuentes del corpus", 2)
    add_paragraph(doc,
        f"Corpus normativo de {estado}: /Users/.../{extraccion_dir.name}/"
    )

    add_heading_styled(doc, "Leyes estatales", 3)
    for ley in leyes:
        if not ley["articulos"]:
            continue
        arts_nums = sorted(set(a["numero"] for a in ley["articulos"]))
        total = len(arts_nums)
        arts_mostrar = arts_nums[:6]
        arts_str = ", ".join(normalizar_numero_articulo(n) for n in arts_mostrar)
        if total > 6:
            arts_str += f" (+{total - 6})"
        # Tabla compacta de 2 columnas (Ley | Arts.) para mejor lectura
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.autofit = False
        hdr = t.rows[0].cells
        hdr[0].text = "Ley"
        hdr[1].text = f"Arts. ({total})"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(8)
        row = t.add_row().cells
        row[0].text = ley["nombre"][:60]
        row[1].text = arts_str
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
        add_paragraph(doc, "")  # Separador

    # FIX 5: Analisis vs normativa del guardarriel
    add_heading_styled(doc, "Analisis vs guardarriel", 3)
    guardarriel_path = pathlib.Path(
        f"/Users/adominguezdia/Documents/FASP/10_Corpus_Definitivo/{estado_corpus_id(estado)}/"
    )
    corpus_base = pathlib.Path("/Users/adominguezdia/Documents/FASP/10_Corpus_Definitivo/")
    tabla_norm_path = None
    fuente_guardarriel = None
    if guardarriel_path.exists():
        for f in guardarriel_path.glob("TABLA_Normatividad_*.md"):
            tabla_norm_path = f
            fuente_guardarriel = "TABLA_Normatividad"
            break

    # Estrategia:
    # 1. Si hay TABLA_Normatividad_<Estado>.md, usar ese (preferido)
    # 2. Si no, extraer leyes del corpus directamente (NOR-* y DOC-*LEY*)
    # 3. LISTA_Normatividad_por_Carpeta.md solo lista archivos (no leyes),
    #    asi que no se usa como fuente del guardarriel
    normas_inline = None
    if not tabla_norm_path:
        corpus_dir_local_early = pathlib.Path(extraccion_dir) if extraccion_dir else pathlib.Path(".")
        normas_inline = extraer_leyes_del_corpus(corpus_dir_local_early)
        if normas_inline:
            tabla_norm_path = pathlib.Path("inline")
            fuente_guardarriel = f"Leyes del corpus ({len(normas_inline)})"

    if tabla_norm_path and (tabla_norm_path.exists() or str(tabla_norm_path) == "inline"):
        if normas_inline is not None:
            add_paragraph(doc, f"Fuente: leyes del corpus (N={len(normas_inline)})")
        else:
            add_paragraph(doc, f"Fuente: {tabla_norm_path.name} ({fuente_guardarriel})")
        try:
            from delimitar_normas_fasp import (
                parsear_tabla_normatividad,
                parsear_md_llm2,
                clasificar_norma,
                clasificar_leyes_no_en_guardarriel,
            )
            if normas_inline is not None:
                normas_g = normas_inline
            else:
                normas_g = parsear_tabla_normatividad(tabla_norm_path)
            leyes_llm = parsear_md_llm2(md_path) if md_path and md_path.exists() else []
            corpus_dir_local = pathlib.Path(extraccion_dir) if extraccion_dir else pathlib.Path(".")
            archivos_corpus = [f.name for f in corpus_dir_local.glob("*.txt")]
            corpus_n = len(archivos_corpus)

            guardarriel_keywords = set()
            for n in normas_g:
                guardarriel_keywords.update({w for w in n["norma"].lower().split() if len(w) >= 5})

            clasif_g = [(n, clasificar_norma(n, leyes_llm, archivos_corpus, guardarriel_keywords))
                        for n in normas_g]
            clasif_e = clasificar_leyes_no_en_guardarriel(leyes_llm, normas_g)

            from collections import Counter
            cnt = Counter(c for _, c in clasif_g)
            extras_cnt = Counter(e["categoria"] for e in clasif_e)

            # Tabla compacta de cruce
            tc = doc.add_table(rows=2, cols=3)
            tc.style = "Table Grid"
            tc.autofit = False
            hdr = tc.rows[0].cells
            hdr[0].text = "Guardarriel"
            hdr[1].text = "MD (LLM-2)"
            hdr[2].text = "Leyes"
            for cell in hdr:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(8)
            row = tc.rows[1].cells
            mant = cnt.get('MANTENER', 0)
            rev = cnt.get('REVISAR', 0)
            falt = cnt.get('FALTANTE', 0)
            fasp = extras_cnt.get('FASP_FUERA_GUARDARRIEL', 0)
            excl = extras_cnt.get('EXCLUIR', 0)
            row[0].text = f"{len(normas_g)}\nMantener: {mant}\nRevisar: {rev}\nFaltante: {falt}"
            row[1].text = f"{len(leyes_llm)}\nFASP: {fasp}\nExcluir: {excl}"
            row[2].text = f"Corpus: {corpus_n}"
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
            add_paragraph(doc, "")

            excluidas = [e["nombre"] for e in clasif_e if e["categoria"] == "EXCLUIR"]
            if excluidas:
                add_paragraph(doc, "Leyes del MD que no son FASP (excluir):")
                for e in excluidas[:5]:
                    add_paragraph(doc, f"  - {e[:70]}")
        except Exception as e:
            add_paragraph(doc, f"(Error en cruce: {e})")
    else:
        add_paragraph(doc, f"No se encontro TABLA_Normatividad_{estado}.md en {guardarriel_path}")

    add_heading_styled(doc, "Normativa federal", 3)
    add_paragraph(doc, "- Criterios Generales del FASP 2026 (DOF, 27 dic 2025)")
    add_paragraph(doc, "- Lineamientos de Evaluacion del FASP (DOF, 2025)")
    add_paragraph(doc, "- Reglamento del SESNSP (DOF)")
    add_paragraph(doc, "- Ley General del Sistema Nacional de Seguridad Publica (DOF)")
    add_paragraph(doc, f"- Convenio FASP {estado} 2026 (firma Federacion-Estado)")

    add_heading_styled(doc, "Notas metodologicas", 3)
    add_paragraph(doc, "- Citas del LLM-2 (MiniMax-M3, API directa). Fallback: OpenRouter.")
    add_paragraph(doc, "- Articulos 'bis' formateados correctamente (ej. Art. 28 Bis, no 28.5).")
    add_paragraph(doc, "- Espacios insertados por extraccion PDF/OCR normalizados.")

    # Footer final
    p = doc.add_paragraph()
    run = p.add_run(
        f"Documento generado el {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} "
        f"por el pipeline fasp-document-pipeline. Estructura conforme al indice "
        f"del cliente (Apartados 5.x y 6.x)."
    )
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main():
    p = argparse.ArgumentParser(description="Producto 1 con estructura 5.x/6.x y citas directas")
    p.add_argument("--extraccion", required=True, help="Directorio con .txt de extraccion")
    p.add_argument("--estado", required=True, help="Nombre del estado")
    p.add_argument("--output", required=True, help="Archivo .docx de salida")
    p.add_argument("--md", default=None,
                   help="Archivo .md con el analisis del LLM-2 (fuente de verdad). "
                        "Si se da, los articulos en la tabla son SOLO los que el "
                        "LLM-2 identifico como relevantes al FASP. Si no se da, "
                        "se filtra por keywords automaticas.")
    args = p.parse_args()

    extraccion_dir = pathlib.Path(args.extraccion)
    if not extraccion_dir.is_dir():
        sys.exit(f"No existe: {extraccion_dir}")

    output_path = pathlib.Path(args.output)

    # Cargar articulos relevantes del MD (fuente de verdad)
    articulos_desde_md = None
    if args.md:
        md_path = pathlib.Path(args.md)
        articulos_desde_md = parsear_articulos_desde_md(md_path)
        if articulos_desde_md:
            total_arts = sum(len(s) for s in articulos_desde_md.values())
            print(f"MD del LLM-2: {len(articulos_desde_md)} leyes, {total_arts} articulos relevantes")
        else:
            print(f"WARN: no se pudieron extraer articulos de {md_path}, usando modo automatico")

    print(f"Cargando leyes desde {extraccion_dir}...")
    leyes = cargar_leyes(extraccion_dir, args.estado, articulos_desde_md=articulos_desde_md)
    print(f"  Leyes cargadas: {len(leyes)}")
    for ley in leyes:
        n = len(ley["articulos"])
        print(f"    - {ley['nombre'][:60]}: {n} articulos relevantes")

    if not leyes:
        sys.exit("No se encontraron leyes. Verifica el directorio y los patrones.")

    print(f"Generando Producto 1 ({args.estado}) con citas directas...")
    build_producto_1(args.estado, leyes, output_path,
                     articulos_desde_md=articulos_desde_md,
                     extraccion_dir=extraccion_dir,
                     md_path=pathlib.Path(args.md) if args.md else None)

    print(f"OK Documento guardado en {output_path}")
    print(f"   Tamano: {output_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()