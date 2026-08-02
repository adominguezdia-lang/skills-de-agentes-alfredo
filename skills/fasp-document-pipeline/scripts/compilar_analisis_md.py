#!/usr/bin/env python3
"""
compilar_analisis_md.py — Convierte un analisis normativo (.md) en Word (.docx).

Uso:
    python3 compilar_analisis_md.py --input apartados_5y6_Queretaro.md \\
        --estado Queretaro --output Sec5y6_Queretaro_2026_NORMATIVO.docx
"""
from __future__ import annotations
import argparse, pathlib, re, sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx no instalado. Ejecuta: pip install python-docx")


# Configuracion por estado (replica el estilo de compile_word.py)
ESTADOS = {
    "Queretaro": {
        "ejecutores": "• SSC-QRO (001)\n• FGE-QRO (002)\n• CIASQ (011)",
        "entes": "• Gobernador del Estado\n• Consejo Estatal de Seguridad\n• Secretaría de Seguridad Ciudadana (SSC-QRO)\n• Fiscalía General del Estado (FGE-QRO)\n• Centro de Información y Análisis (CIASQ)",
    },
    "Michoacan": {
        "ejecutores": "• SSP (001)\n• FGE Mich (002)\n• SESESP",
        "entes": "• Gobernador del Estado\n• Consejo Estatal de Seguridad Pública\n• Secretaría de Seguridad Pública (SSP)\n• Fiscalía General del Estado\n• SESESP\n• Secretaría de Finanzas y Administración (SFA)\n• Auditoría Superior de Michoacán (ASM)",
    },
    "Hidalgo": {"ejecutores": "• SSPH (001)\n• PGJEH (002)", "entes": "• Gobernador del Estado\n• Consejo Estatal de Seguridad"},
    "EdoMexico": {"ejecutores": "• SSEM (001)\n• FGJEM (002)", "entes": "• Gobernadora del Estado\n• Consejo Estatal de Seguridad"},
}


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x44, 0x80)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2D, 0x5E, 0x8E)
    else:
        run.font.size = Pt(12)
        run.font.bold = True


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True


def add_table_from_markdown(doc: Document, md_table: str):
    """Convierte una tabla Markdown a una tabla Word."""
    lines = [l for l in md_table.strip().split("\n") if l.strip()]
    if len(lines) < 2:
        return
    # Header
    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    # Data (skip separator line |---|---|)
    data = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        data.append(cells)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Background color for header (dark blue)
    for cell in hdr_cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1A4480")
        tc_pr.append(shd)
    # Data rows
    for row in data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(cells):
                cells[i].text = ""
                p = cells[i].paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)


def md_to_docx(md_text: str, estado: str) -> Document:
    doc = Document()
    # Margenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Encabezado
    info = ESTADOS.get(estado, {})
    add_heading(doc, f"Apartados 5 y 6 — Análisis Normativo", 1)
    add_heading(doc, f"Evaluación Estratégica de Coordinación al FASP", 2)
    add_paragraph(doc, f"Estado: {estado}", bold=True)
    add_paragraph(doc, f"Programa: Fondo de Aportaciones para la Seguridad Pública (FASP)")
    add_paragraph(doc, f"Ejercicio Fiscal: 2026")
    add_paragraph(doc, f"Fecha de elaboracion: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    # Ejecutores
    if info.get("ejecutores"):
        add_heading(doc, "Entes ejecutores", 3)
        for line in info["ejecutores"].split("\n"):
            add_paragraph(doc, line)
    # Entes
    if info.get("entes"):
        add_heading(doc, "Entes involucrados en la coordinacion", 3)
        for line in info["entes"].split("\n"):
            add_paragraph(doc, line)
    doc.add_paragraph()

    # Procesar el Markdown
    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        # Detectar tabla Markdown (linea con | y ---)
        if line.strip().startswith("|") and "|" in line:
            # Acumular tabla
            in_table = True
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            # Procesar tabla acumulada
            if table_buffer:
                add_table_from_markdown(doc, "\n".join(table_buffer))
                table_buffer = []
            in_table = False
            doc.add_paragraph()

        # Detectar headings
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            # Lista
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip():
            # Parrafo normal
            add_paragraph(doc, line.strip())
        i += 1

    # Si quedo tabla pendiente
    if table_buffer:
        add_table_from_markdown(doc, "\n".join(table_buffer))

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"\nDocumento generado el {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} por el pipeline fasp-document-pipeline.")
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    return doc


def main():
    p = argparse.ArgumentParser(description="Convertir analisis normativo .md a .docx")
    p.add_argument("--input", required=True, help="Archivo .md de entrada")
    p.add_argument("--estado", required=True, help="Nombre del estado")
    p.add_argument("--output", required=True, help="Archivo .docx de salida")
    args = p.parse_args()

    md_path = pathlib.Path(args.input)
    if not md_path.exists():
        sys.exit(f"No existe: {md_path}")

    output_path = pathlib.Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"Leyendo {md_path}...")
    md_text = md_path.read_text(encoding="utf-8")

    print(f"Generando documento Word para {args.estado}...")
    doc = md_to_docx(md_text, args.estado)
    doc.save(str(output_path))

    print(f"OK Documento guardado en {output_path}")
    print(f"   Tamano: {output_path.stat().st_size:,} bytes")
    print(f"   Para abrirlo: open {output_path}")


if __name__ == "__main__":
    main()